from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.shared import Inches
from tkinter import *

tk = Tk()
tk.title('导出word')
tk.geometry('350x350')
#  输入框1
now_nub1 = Label(tk, text='1、学校：')
now_nub1.grid(row=1, column=1, sticky='E')
now_bok1 = Entry(tk, bd=5)
now_bok1.grid(row=1, column=2, sticky='NW')

#  输入框2
now_nub2 = Label(tk, text='2、班级：')
now_nub2.grid(row=2, column=1, sticky='E')
now_bok2 = Entry(tk, bd=5)
now_bok2.grid(row=2, column=2, sticky='NW')
#  输入框3
now_nub3 = Label(tk, text='3、姓名：')
now_nub3.grid(row=3, column=1, sticky='E')
now_bok3 = Entry(tk, bd=5)
now_bok3.grid(row=3, column=2, sticky='NW')

now_nub4 = Label(tk, text='4、座号：')
now_nub4.grid(row=4, column=1, sticky='E')
now_bok4 = Entry(tk, bd=5)
now_bok4.grid(row=4, column=2, sticky='NW')


def word():
    document = Document()

    default_section = document.sections[0]
    default_section.top_margin = Cm(1.27)
    default_section.right_margin = Cm(1.27)
    default_section.bottom_margin = Cm(1.27)
    default_section.left_margin = Cm(1.27)

    school = now_bok1.get()
    cl = now_bok2.get()
    name = now_bok3.get()
    number = now_bok4.get()

    # 增加表格
    table = document.add_table(rows=8, cols=3, style='Table Grid')

    table.style.paragraph_format.line_spacing = 1.4

    for index in range(8):
        hdr_cells = table.rows[index].cells
        hdr_cells[0].text = school + '\n班级：' + cl + '\n姓名：' + name + '\n座号：' + number
        hdr_cells[1].text = school + '\n班级：' + cl + '\n姓名：' + name + '\n座号：' + number
        hdr_cells[2].text = school + '\n班级：' + cl + '\n姓名：' + name + '\n座号：' + number

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(12)
                    font.bold = True
                    font.name = u'楷体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

    # 保存文件
    document.save(cl+name+number+'.docx')

AnNiu = Button(tk, text='提交', fg='blue', bd=2, width=10, command=word)
AnNiu.grid(row=5, column=2, sticky='NW')

tk.mainloop()


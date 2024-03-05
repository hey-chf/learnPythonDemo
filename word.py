from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.shared import Inches

# 打开文档
document = Document()

default_section = document.sections[0]
default_section.top_margin = Cm(1.27)
default_section.right_margin = Cm(1.27)
default_section.bottom_margin = Cm(1.27)
default_section.left_margin = Cm(1.27)

paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.line_spacing = 1.4
school = '仓山区金港湾实验学校'
cl = '二年（2）班'
name = '陈辉凤'
number = '52号'

# 增加表格
table = document.add_table(rows=8, cols=3, style='Table Grid')

for index in range(8):
    hdr_cells = table.rows[index].cells
    hdr_cells[0].text = school + '\n班级：' + cl + '\n姓名：' + name + '\n座号：' + number
    hdr_cells[1].text = school + '\n班级：' + cl + '\n姓名：' + name + '\n座号：' + number
    hdr_cells[2].text = school + '\n班级：' + cl + '\n姓名：' + name + '\n座号：' + number

for row in table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(12)
                font.bold = True
                font.name = u'楷体'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

# 保存文件
document.save('demo2.docx')